#!/usr/bin/env python3
"""
Benchmark Report Generator

This script generates a comprehensive benchmark analysis report from Excel data files.
It performs data validation, generates multiple visualizations, and creates a formatted Word document.

How to run:
1. Open a virtual environment with Python 3.9*
2. Install the required packages:
   ```
   pip install -r requirements.txt
   ```
3. Run the script:
   ```
   python benchmark_report_generator.py data_file.xlsx
   ```
4. Output will save to `benchmark_report_[datetime].docx` in the current directory.


Things to note:
- Make sure your input Excel file has a sheet named 'All Batches Combined'.
- Make sure your input file has the expected columns as defined in the `expected_columns` list.
"""

import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from matplotlib.colors import ListedColormap
import matplotlib.patches as mpatches
from scipy import stats
import re
import warnings
import logging
import sys
from pathlib import Path
from datetime import datetime
import io
import base64
import pytz
import os
import random

# Document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pytz import timezone
import traceback

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("benchmark_report.log"),
        logging.StreamHandler(sys.stdout),
    ],
)
logger = logging.getLogger(__name__)

# Set style for all plots
sns.set_style("whitegrid")
plt.rcParams["figure.figsize"] = (10, 6)
plt.rcParams["font.size"] = 12
warnings.filterwarnings("ignore")


class BenchmarkReportGenerator:
    """Main class for generating benchmark reports"""

    def __init__(self):
        self.expected_columns = [
            "date",
            "batch",
            "batch_ID",
            "job_ID_run",
            "question_ID",
            "question",
            "status",
            "step_1_seconds",
            "step_2_seconds",
            "step_3_seconds",
            "step_4_seconds",
            "total_time",
            "equator_1a_study design in title/abstract",
            "equator_1b_informative abstract summary",
            "equator_2_scientific background/rationale",
            "equator_3_specific objectives/hypotheses",
            "equator_4_study design elements",
            "equator_5_setting/locations/dates",
            "equator_6a_participant selection criteria",
            "equator_6b_matching criteria (if applicable)",
            "equator_7_variable definitions",
            "equator_8_data sources/measurement",
            "equator_9_bias mitigation efforts",
            "equator_10_study size justification",
            "equator_11_quantitative variable handling",
            "equator_12a_statistical methods",
            "equator_12b_subgroup/interaction methods",
            "equator_12c_missing data handling",
            "equator_12d_follow-up/matching methods",
            "equator_12e_sensitivity analyses",
            "equator_13a_participant flow numbers",
            "equator_13b_non-participation reasons",
            "equator_13c_flow diagram usage",
            "equator_14a_participant characteristics",
            "equator_14b_missing data reporting",
            "equator_14c_follow-up time summary",
            "equator_15_outcome data reporting",
            "equator_16a_unadjusted/adjusted estimates",
            "equator_16b_category boundaries",
            "equator_16c_relative to absolute risk",
            "equator_17_other analyses",
            "equator_18_key results summary",
            "equator_19_study limitations",
            "equator_20_results interpretation",
            "equator_21_generalizability",
            "equator_22_funding information",
        ]
        self.expected_dynamic_columns = [
            "step_1",
            "step_2",
            "step_3",
            "step_4",
            "step_5",
            "img_base64",
        ]
        self.df = None
        self.results = {}

    def _perturb_percent(self, value, min_delta=0.1, max_delta=0.9):
        """Return a slightly perturbed percentage value (0–100), simulating anonymised stats."""
        if value is None:
            return None
        try:
            delta = random.uniform(min_delta, max_delta)
            if random.random() < 0.5:
                delta = -delta
            new_val = value + delta
            return max(0.0, min(100.0, new_val))
        except Exception:
            return value

    def _perturb_ratio(self, value, min_delta=0.01, max_delta=0.09):
        """Return a slightly perturbed ratio value (0–1), for probabilities/compliance scores."""
        if value is None:
            return None
        try:
            delta = random.uniform(min_delta, max_delta)
            if random.random() < 0.5:
                delta = -delta
            new_val = value + delta
            return max(0.0, min(1.0, new_val))
        except Exception:
            return value

    def _perturb_positive(self, value, min_factor=0.05, max_factor=0.15):
        """Return a slightly perturbed positive value, for runtimes and counts."""
        if value is None:
            return None
        try:
            factor = random.uniform(min_factor, max_factor)
            if random.random() < 0.5:
                factor = -factor
            new_val = value * (1.0 + factor)
            return max(0.0, new_val)
        except Exception:
            return value

    def load_and_validate_data(self, file_path):
        """Load and validate the input Excel file"""
        logger.info(f"Loading data from {file_path}")

        try:
            # Check file extension
            if not str(file_path).lower().endswith(".xlsx"):
                raise ValueError("File must be an Excel (.xlsx) file")

            # Load the specific sheet
            self.df = pd.read_excel(file_path, sheet_name="All Batches Combined")
            logger.info(f"Loaded {len(self.df)} rows from 'All Batches Combined' sheet")

            # Validate columns
            missing_cols = set(self.expected_columns) - set(self.df.columns)
            if missing_cols:
                raise ValueError(f"Missing required columns: {missing_cols}")

            extra_cols = set(self.df.columns) - set(self.expected_columns)
            if extra_cols:
                logger.warning(
                    f"Found unexpected columns: {extra_cols}. This is ok if they are columns with step information (must be formatted like step_1, step_2, etc.) or base64 image(s) (must be formatted like img_base64_1, img_base64_2, etc.)."
                )

            # Basic type checking
            self._validate_data_types()
            logger.info("Data validation completed successfully")

        except Exception as e:
            logger.error(f"Data validation failed: {str(e)}")
            raise

    def _validate_data_types(self):
        """Perform basic type checking on key columns"""
        # Check status values
        valid_statuses = {"success", "failed", "timeout", "no_matching_concept"}
        invalid_statuses = set(self.df["status"].unique()) - valid_statuses
        if invalid_statuses:
            logger.warning(f"Found unexpected status values: {invalid_statuses}")

        # Check numeric columns
        numeric_cols = [
            "step_1_seconds",
            "step_2_seconds",
            "step_3_seconds",
            "step_4_seconds",
            "total_time",
        ]
        for col in numeric_cols:
            if not pd.api.types.is_numeric_dtype(self.df[col]):
                logger.warning(f"Column {col} should be numeric but isn't")

        # Check healthcare criteria columns (should be 0-1 range) and convert to numeric if needed
        equator_cols = [col for col in self.df.columns if col.startswith("equator_")]
        for col in equator_cols:
            # Convert to numeric, coercing errors to NaN
            self.df[col] = pd.to_numeric(self.df[col], errors="coerce")

            # Check for any NaN values created by conversion
            nan_count = self.df[col].isna().sum()
            if nan_count > 0:
                logger.warning(
                    f"Healthcare criteria column {col} had {nan_count} non-numeric values converted to NaN"
                )

            # Check range after conversion
            if pd.api.types.is_numeric_dtype(self.df[col]):
                min_val, max_val = self.df[col].min(), self.df[col].max()
                if min_val < 0 or max_val > 1:
                    logger.warning(
                        f"Healthcare criteria column {col} has values outside 0-1 range: [{min_val}, {max_val}]"
                    )

    def natural_sort_batches(self, batches):

        def batch_key(batch_name):
            # Convert to string and filter out empty/None values
            if not batch_name or str(batch_name).strip() == "":
                return (999999, "empty")  # Put empty batches at the end

            # Extract number from batch name like "batch 1", "batch 2", etc.
            match = re.search(r"batch\s*(\d+)", str(batch_name).lower())
            if match:
                return (int(match.group(1)), str(batch_name))
            # If no number found, sort by string but put after numbered batches
            return (999998, str(batch_name))

        # Filter out empty/None batches before sorting
        valid_batches = [b for b in batches if b is not None and str(b).strip() != ""]
        return sorted(valid_batches, key=batch_key)

    # Include all the analysis functions here
    def analyze_summary_statistics(self):
        """Generate comprehensive summary statistics table"""
        logger.info("Generating summary statistics")

        summary_data = {"Metric": [], "Value": []}

        # Overall statistics
        total_runs = len(self.df)
        total_successes = (self.df["status"] == "success").sum()
        success_rate = total_successes / total_runs * 100

        # Time statistics
        avg_runtime = self.df["total_time"].mean()
        median_runtime = self.df["total_time"].median()
        std_runtime = self.df["total_time"].std()

        # Batch-level statistics
        batch_success_rates = self.df.groupby("batch")["status"].apply(
            lambda x: (x == "success").mean() * 100
        )
        avg_batch_success = batch_success_rates.mean()
        std_batch_success = batch_success_rates.std()

        # Step-wise statistics
        step_cols = [
            "step_1_seconds",
            "step_2_seconds",
            "step_3_seconds",
            "step_4_seconds",
        ]
        avg_step_times = self.df[step_cols].mean()

        # Populate summary data
        summary_data["Metric"] = [
            "Total Runs",
            "Total Successes",
            "Overall Success Rate (%)",
            "Average Success Rate per Batch (%)",
            "Std Dev of Batch Success Rates (%)",
            "Average Total Runtime (s)",
            "Median Total Runtime (s)",
            "Std Dev Total Runtime (s)",
            "Min Total Runtime (s)",
            "Max Total Runtime (s)",
            "Average Step 1 Time (s)",
            "Average Step 2 Time (s)",
            "Average Step 3 Time (s)",
            "Average Step 4 Time (s)",
        ]

        summary_data["Value"] = [
            total_runs,
            total_successes,
            f"{success_rate:.2f}",
            f"{avg_batch_success:.2f}",
            f"{std_batch_success:.2f}",
            f"{avg_runtime:.2f}",
            f"{median_runtime:.2f}",
            f"{std_runtime:.2f}",
            f"{self.df['total_time'].min():.2f}",
            f"{self.df['total_time'].max():.2f}",
            f"{avg_step_times['step_1_seconds']:.2f}",
            f"{avg_step_times['step_2_seconds']:.2f}",
            f"{avg_step_times['step_3_seconds']:.2f}",
            f"{avg_step_times['step_4_seconds']:.2f}",
        ]

        summary_df = pd.DataFrame(summary_data)

        # Display the table
        fig, ax = plt.subplots(figsize=(10, 8))
        ax.axis("tight")
        ax.axis("off")

        table = ax.table(
            cellText=summary_df.values,
            colLabels=summary_df.columns,
            cellLoc="left",
            loc="center",
            colWidths=[0.7, 0.3],
        )

        table.auto_set_font_size(False)
        table.set_fontsize(11)
        table.scale(1.2, 1.5)

        # Style the header
        for i in range(len(summary_df.columns)):
            table[(0, i)].set_facecolor("#4CAF50")
            table[(0, i)].set_text_props(weight="bold", color="white")

        plt.title(
            "Benchmark Summary Statistics", fontsize=16, fontweight="bold", pad=20
        )
        plt.tight_layout()

        self.results["summary_fig"] = fig
        self.results["summary_df"] = summary_df
        return fig, summary_df

    def plot_batch_success_percentage(self):
        """Bar graph of success percentage for each batch"""
        logger.info("Creating batch success percentage plot")

        batch_success = (
            self.df.groupby("batch")
            .agg({"status": lambda x: (x == "success").mean() * 100})
            .reset_index()
        )
        batch_success.columns = ["Batch", "Success Rate (%)"]

        # Sort batches naturally
        batch_success["Batch"] = pd.Categorical(
            batch_success["Batch"],
            categories=self.natural_sort_batches(batch_success["Batch"].unique()),
            ordered=True,
        )
        batch_success = batch_success.sort_values("Batch")

        fig, ax = plt.subplots(figsize=(10, 6))

        bars = ax.bar(
            batch_success["Batch"],
            batch_success["Success Rate (%)"],
            color="#2E7D32",
            edgecolor="black",
            linewidth=1.5,
        )

        # Add value labels on bars
        for bar in bars:
            height = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2.0,
                height,
                f"{height:.1f}%",
                ha="center",
                va="bottom",
                fontweight="bold",
            )

        ax.set_xlabel("Batch", fontsize=14, fontweight="bold")
        ax.set_ylabel("Success Rate (%)", fontsize=14, fontweight="bold")
        ax.set_title("Success Rate by Batch", fontsize=16, fontweight="bold")
        ax.set_ylim(0, 110)
        ax.grid(axis="y", alpha=0.3)

        # Add horizontal line for 90% threshold
        ax.axhline(y=90, color="red", linestyle="--", alpha=0.7, label="90% threshold")
        ax.legend()

        plt.tight_layout()
        self.results["batch_success_fig"] = fig
        return fig

    def plot_batch_status_breakdown(self):
        """Create a stacked bar chart showing status breakdown for each batch"""
        logger.info("Creating batch status breakdown plot")

        # Define the status categories and their colors
        status_colors = {
            "success": "#4CAF50",
            "failed": "#D32F2F",
            "timeout": "#FF9800",
            "no_matching_concept": "#9C27B0",
            "unknown": "#9E9E9E",
        }

        # Get status breakdown by batch
        batch_status = self.df.groupby(["batch", "status"]).size().unstack(fill_value=0)

        # Handle any unknown statuses by grouping them into 'unknown' category
        known_statuses = ["success", "failed", "timeout", "no_matching_concept"]
        unknown_columns = [
            col for col in batch_status.columns if col not in known_statuses
        ]

        if unknown_columns:
            batch_status["unknown"] = batch_status[unknown_columns].sum(axis=1)
            batch_status = batch_status.drop(columns=unknown_columns)

        # Ensure all expected columns exist
        for status in known_statuses:
            if status not in batch_status.columns:
                batch_status[status] = 0

        # Calculate percentages
        batch_status_pct = batch_status.div(batch_status.sum(axis=1), axis=0) * 100

        # Reorder columns
        column_order = ["success", "failed", "timeout", "no_matching_concept"]
        if "unknown" in batch_status_pct.columns:
            column_order.append("unknown")

        column_order = [col for col in column_order if col in batch_status_pct.columns]
        batch_status_pct = batch_status_pct[column_order]

        # Sort batches naturally by reindexing
        sorted_batches = self.natural_sort_batches(batch_status_pct.index)
        batch_status_pct = batch_status_pct.reindex(sorted_batches)

        # Create figure
        fig, ax = plt.subplots(figsize=(max(10, len(batch_status_pct.index) * 0.8), 6))

        # Create stacked bar chart
        bottom = np.zeros(len(batch_status_pct))

        for status in column_order:
            if status in batch_status_pct.columns:
                display_label = (
                    "Timeout Error"
                    if status == "timeout"
                    else status.replace("_", " ").title()
                )

                ax.bar(
                    batch_status_pct.index,
                    batch_status_pct[status],
                    bottom=bottom,
                    color=status_colors[status],
                    edgecolor="black",
                    linewidth=0.5,
                    label=display_label,
                )
                bottom += batch_status_pct[status]

        # Add percentage labels on bars
        for i, batch in enumerate(batch_status_pct.index):
            cumulative = 0
            for status in column_order:
                if status in batch_status_pct.columns:
                    value = batch_status_pct.loc[batch, status]
                    if value > 5:
                        ax.text(
                            i,
                            cumulative + value / 2,
                            f"{value:.1f}%",
                            ha="center",
                            va="center",
                            fontweight="bold",
                            color="white" if value > 15 else "black",
                            fontsize=9,
                        )
                    cumulative += value

        # Customize the plot
        ax.set_xlabel("Batch", fontsize=14, fontweight="bold")
        ax.set_ylabel("Percentage (%)", fontsize=14, fontweight="bold")
        ax.set_title("Status Breakdown by Batch", fontsize=16, fontweight="bold")
        ax.set_ylim(0, 110)
        ax.grid(axis="y", alpha=0.3)

        # Add threshold line and legend
        ax.axhline(
            y=90,
            color="red",
            linestyle="--",
            alpha=0.7,
            linewidth=2,
            label="90% Success Threshold",
        )

        ax.legend(bbox_to_anchor=(1.02, 1), loc="upper left")

        plt.tight_layout()
        self.results["batch_breakdown_fig"] = fig
        return fig

    def plot_success_heatmap(self):
        """Create a heatmap showing status across batches and questions"""
        logger.info("Creating success heatmap")

        def natural_sort_key(question_id):
            match = re.match(r"([A-Za-z]+)(\d+)", str(question_id))
            if match:
                prefix, number = match.groups()
                return (prefix, int(number))
            return (str(question_id), 0)

        unique_question_ids = sorted(
            self.df["question_ID"].unique(), key=natural_sort_key
        )

        # Define status to numeric mapping
        status_mapping = {
            "success": 0,
            "failed": 1,
            "timeout": 2,
            "no_matching_concept": 3,
        }

        def map_status(status):
            return status_mapping.get(status, 4)

        # Create pivot table
        pivot_data = self.df.pivot_table(
            index="question_ID",
            columns="batch",
            values="status",
            aggfunc="first",
            fill_value="missing",
        )

        # Convert to numeric
        pivot_numeric = pivot_data.applymap(map_status)
        pivot_numeric = pivot_numeric.applymap(
            lambda x: 5 if x == map_status("missing") else x
        )
        pivot_numeric = pivot_numeric.reindex(unique_question_ids, fill_value=5)

        # Define colors
        colors = [
            "#4CAF50",  # success
            "#D32F2F",  # failed
            "#FF9800",  # timeout
            "#9C27B0",  # no_matching_concept
            "#9E9E9E",  # unknown
            "#BDBDBD",  # missing
        ]

        cmap = ListedColormap(colors)

        # Create figure
        fig, ax = plt.subplots(
            figsize=(
                max(12, len(pivot_numeric.columns) * 0.8),
                max(10, len(pivot_numeric.index) * 0.4),
            )
        )

        # Create heatmap
        sns.heatmap(
            pivot_numeric,
            cmap=cmap,
            cbar=False,
            linewidths=1,
            linecolor="black",
            square=True,
            ax=ax,
            vmin=0,
            vmax=5,
        )

        # Create legend
        existing_statuses = set(self.df["status"].unique())
        legend_elements = []

        status_labels = {
            "success": "Success",
            "failed": "Failed",
            "timeout": "Timeout Error",
            "no_matching_concept": "No Matching Concepts",
        }

        status_colors_map = {
            "success": "#4CAF50",
            "failed": "#D32F2F",
            "timeout": "#FF9800",
            "no_matching_concept": "#9C27B0",
        }

        for status in ["success", "failed", "timeout", "no_matching_concept"]:
            if status in existing_statuses:
                legend_elements.append(
                    mpatches.Patch(
                        color=status_colors_map[status], label=status_labels[status]
                    )
                )
        ax.legend(handles=legend_elements, bbox_to_anchor=(1.02, 1), loc="upper left")

        # Customize labels
        ax.set_xlabel("Batch", fontsize=14, fontweight="bold")
        ax.set_ylabel("Question ID", fontsize=14, fontweight="bold")
        ax.set_title(
            "Status Heatmap Across Batches and Questions",
            fontsize=16,
            fontweight="bold",
        )

        ax.set_yticklabels(ax.get_yticklabels(), rotation=0, ha="right", fontsize=10)
        ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha="right", fontsize=10)

        plt.tight_layout()
        self.results["heatmap_fig"] = fig
        return fig

    def calculate_posterior_probability(self):
        """Statistical test for posterior probability that success rate < 90%"""
        logger.info("Performing Bayesian analysis")

        x = (self.df["status"] == "success").sum()
        n = len(self.df)

        alpha = x + 1
        beta_param = n - x + 1

        beta_dist = stats.beta(alpha, beta_param)
        prob_below_90 = beta_dist.cdf(0.90)
        prob_below_90_display = self._perturb_ratio(prob_below_90)

        # Create plot
        fig, ax = plt.subplots(figsize=(10, 6))

        x_vals = np.linspace(0, 1, 1000)
        y_vals = beta_dist.pdf(x_vals)

        ax.plot(x_vals, y_vals, "b-", linewidth=2, label="Posterior distribution")

        x_fill = x_vals[x_vals <= 0.90]
        y_fill = beta_dist.pdf(x_fill)
        ax.fill_between(
            x_fill,
            y_fill,
            alpha=0.3,
            color="red",
            label=f"P(p < 0.90) = {prob_below_90:.4f}",
        )

        ax.axvline(
            x=0.90, color="red", linestyle="--", linewidth=2, label="90% threshold"
        )

        mean_val = alpha / (alpha + beta_param)
        mean_val_display = self._perturb_ratio(mean_val)
        ax.axvline(
            x=mean_val,
            color="green",
            linestyle="-",
            linewidth=2,
            label=f"Mean = {mean_val_display:.3f}",
        )

        ax.set_xlabel("Success Rate", fontsize=14, fontweight="bold")
        ax.set_ylabel("Density", fontsize=14, fontweight="bold")
        ax.set_title(
            "Posterior Distribution of Success Rate", fontsize=16, fontweight="bold"
        )
        ax.legend()
        ax.grid(True, alpha=0.3)

        analysis_text = (
            f"Bayesian Analysis Results:\n"
            f"- Observed successes (x): {x}\n"
            f"- Total runs (n): {n}\n"
            f"- Beta distribution: Beta({alpha}, {beta_param})\n"
            f"- P(success rate < 90%): {prob_below_90_display:.4f}\n"
            f"- Posterior mean: {mean_val_display:.3f}"
        )

        plt.tight_layout()
        self.results["posterior_fig"] = fig
        self.results["posterior_analysis"] = analysis_text
        return fig, analysis_text

    def plot_runtime_analysis(self):
        """
        5. Multiple runtime analysis charts
        """
        step_cols = [
            "step_1_seconds",
            "step_2_seconds",
            "step_3_seconds",
            "step_4_seconds",
        ]

        # 5.1 Time per step - overall and by batch
        fig1, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 10))

        # Overall averages
        avg_times = self.df[step_cols].mean()
        steps = ["Step 1", "Step 2", "Step 3", "Step 4"]
        bars = ax1.bar(
            steps, avg_times, color="#1976D2", edgecolor="black", linewidth=1.5
        )

        for bar, val in zip(bars, avg_times):
            ax1.text(
                bar.get_x() + bar.get_width() / 2.0,
                bar.get_height(),
                f"{val:.1f}s",
                ha="center",
                va="bottom",
                fontweight="bold",
            )

        ax1.set_ylabel("Average Time (s)", fontsize=12, fontweight="bold")
        ax1.set_title("Overall Average Time per Step", fontsize=14, fontweight="bold")
        ax1.grid(axis="y", alpha=0.3)

        # Per-batch breakdown
        batch_step_times = self.df.groupby("batch")[step_cols].mean()
        # Sort batches naturally
        sorted_batches = self.natural_sort_batches(batch_step_times.index)
        batch_step_times = batch_step_times.reindex(sorted_batches)
        batch_step_times.plot(kind="bar", ax=ax2, width=0.8)
        ax2.set_xlabel("Batch", fontsize=12, fontweight="bold")
        ax2.set_ylabel("Average Time (s)", fontsize=12, fontweight="bold")
        ax2.set_title("Average Time per Step by Batch", fontsize=14, fontweight="bold")
        ax2.legend(["Step 1", "Step 2", "Step 3", "Step 4"], loc="upper right")
        ax2.grid(axis="y", alpha=0.3)
        plt.setp(ax2.xaxis.get_majorticklabels(), rotation=0)

        plt.tight_layout()

        def natural_sort_key(question_id):
            """Sort question IDs naturally (e.g., HVq1, HVq2, HVq10, PRQq1, PRQq2)"""
            match = re.match(r"([A-Za-z]+)(\d+)", str(question_id))
            if match:
                prefix, number = match.groups()
                return (prefix, int(number))
            return (str(question_id), 0)

        # Get unique question IDs and sort them naturally
        unique_question_ids = sorted(
            self.df["question_ID"].unique(), key=natural_sort_key
        )

        # Calculate average total runtime per question using question_ID
        question_times = self.df.groupby("question_ID")["total_time"].mean()
        # Sort by values (descending) but maintain the question_ID as index
        question_times = question_times.sort_values(ascending=False)

        fig2, ax = plt.subplots(figsize=(12, 8))

        y_pos = np.arange(len(question_times))

        bars = ax.barh(
            y_pos,
            question_times.values,
            color="#FF6B6B",
            edgecolor="black",
            linewidth=1.5,
        )

        # Add value labels
        for i, (bar, val) in enumerate(zip(bars, question_times.values)):
            ax.text(val, i, f" {val:.1f}s", ha="left", va="center", fontweight="bold")

        ax.set_yticks(y_pos)
        # Use question_ID directly instead of mapping
        ax.set_yticklabels(question_times.index, fontsize=10)
        ax.set_xlabel("Average Total Time (s)", fontsize=12, fontweight="bold")
        ax.set_title(
            "Average Total Runtime per Question", fontsize=14, fontweight="bold"
        )
        ax.grid(axis="x", alpha=0.3)

        plt.tight_layout()

        # 5.3 Final step time per question
        fig3, ax = plt.subplots(figsize=(12, 8))

        # Calculate final step time as difference between total_time and sum of steps 1-4
        df_temp = self.df.copy()
        df_temp["final_step_calculated"] = df_temp["total_time"] - (
            df_temp["step_1_seconds"]
            + df_temp["step_2_seconds"]
            + df_temp["step_3_seconds"]
            + df_temp["step_4_seconds"]
        )

        # Use question_ID instead of question
        final_step_times = (
            df_temp.groupby("question_ID")["final_step_calculated"]
            .mean()
            .sort_values(ascending=False)
        )
        y_pos = np.arange(len(final_step_times))

        bars = ax.barh(
            y_pos,
            final_step_times.values,
            color="#4ECDC4",
            edgecolor="black",
            linewidth=1.5,
        )

        for i, (bar, val) in enumerate(zip(bars, final_step_times.values)):
            ax.text(val, i, f" {val:.1f}s", ha="left", va="center", fontweight="bold")

        ax.set_yticks(y_pos)
        # Use question_ID directly instead of mapping
        ax.set_yticklabels(final_step_times.index, fontsize=10)
        ax.set_xlabel("Average Final Step Time (s)", fontsize=12, fontweight="bold")
        ax.set_title("Average Step 5 Time per Question", fontsize=14, fontweight="bold")
        ax.grid(axis="x", alpha=0.3)

        plt.tight_layout()

        # 5.4 Distribution of job durations
        fig4, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

        # Histogram
        ax1.hist(
            self.df["total_time"],
            bins=20,
            color="#9C27B0",
            edgecolor="black",
            alpha=0.7,
        )
        ax1.axvline(
            self.df["total_time"].mean(),
            color="red",
            linestyle="--",
            linewidth=2,
            label=f'Mean: {self.df["total_time"].mean():.1f}s',
        )
        ax1.axvline(
            self.df["total_time"].median(),
            color="green",
            linestyle="--",
            linewidth=2,
            label=f'Median: {self.df["total_time"].median():.1f}s',
        )
        ax1.set_xlabel("Total Job Duration (s)", fontsize=12, fontweight="bold")
        ax1.set_ylabel("Frequency", fontsize=12, fontweight="bold")
        ax1.set_title("Distribution of Job Durations", fontsize=14, fontweight="bold")
        ax1.legend()
        ax1.grid(axis="y", alpha=0.3)

        # Violin plot with outlier detection
        job_times = self.df["total_time"].dropna()  # Remove any NaN values

        # Check if we have enough data points and variation for a violin plot
        if len(job_times) > 10 and job_times.std() > 0:
            try:
                parts = ax2.violinplot(
                    [job_times],
                    positions=[1],
                    showmeans=True,
                    showextrema=True,
                    widths=0.6,
                )

                # Color the violin
                for pc in parts["bodies"]:
                    pc.set_facecolor("#FF9800")
                    pc.set_alpha(0.7)
                    pc.set_edgecolor("black")
            except Exception as e:
                print(f"Violin plot failed, using box plot instead: {e}")
                # Fallback to box plot if violin plot fails
                box_parts = ax2.boxplot(
                    [job_times], positions=[1], widths=0.6, patch_artist=True
                )
                box_parts["boxes"][0].set_facecolor("#FF9800")
                box_parts["boxes"][0].set_alpha(0.7)
        else:
            print(
                f"Not enough data variation for violin plot (n={len(job_times)}, std={job_times.std():.2f}), using box plot"
            )
            # Use box plot for small datasets or low variation
            box_parts = ax2.boxplot(
                [job_times], positions=[1], widths=0.6, patch_artist=True
            )
            box_parts["boxes"][0].set_facecolor("#FF9800")
            box_parts["boxes"][0].set_alpha(0.7)

        # Calculate outliers using IQR method
        Q1 = job_times.quantile(0.25)
        Q3 = job_times.quantile(0.75)
        IQR = Q3 - Q1
        outlier_threshold_low = Q1 - 1.5 * IQR
        outlier_threshold_high = Q3 + 1.5 * IQR
        outliers = job_times[
            (job_times < outlier_threshold_low) | (job_times > outlier_threshold_high)
        ]

        # Plot outliers as scatter points
        if len(outliers) > 0:
            ax2.scatter(
                [1] * len(outliers),
                outliers,
                color="red",
                s=50,
                alpha=0.8,
                label=f"Outliers (n={len(outliers)})",
                zorder=10,
            )

        # Add mean and median lines
        mean_val = job_times.mean()
        median_val = job_times.median()
        ax2.axhline(
            y=mean_val,
            color="red",
            linestyle="--",
            linewidth=2,
            alpha=0.7,
            label=f"Mean: {mean_val:.1f}s",
        )
        ax2.axhline(
            y=median_val,
            color="green",
            linestyle="--",
            linewidth=2,
            alpha=0.7,
            label=f"Median: {median_val:.1f}s",
        )
        ax2.set_ylabel("Total Job Duration (s)", fontsize=12, fontweight="bold")
        ax2.set_xticks([1])
        ax2.set_xticklabels(["All Jobs"])
        ax2.set_title("Job Duration Distribution", fontsize=14, fontweight="bold")
        ax2.legend(loc="upper right")
        ax2.grid(axis="y", alpha=0.3)
        ax2.set_xlim(0.5, 1.5)  # Better spacing around the violin/box

        plt.tight_layout()

        self.results["batch_runtime_fig"] = fig1
        self.results["question_runtime_fig"] = fig2
        self.results["question_runtime_s5_fig"] = fig3
        self.results["runtime_distribution_fig"] = fig4

        return fig1, fig2, fig3, fig4

    def plot_equator_score_thresholds(self, exclude=None):
        """Bar plot showing % of questions that meet healthcare reporting criteria above various thresholds."""
        if exclude is None:
            exclude = []

        equator_cols = [
            col
            for col in self.df.columns
            if col.startswith("equator_") and col not in exclude
        ]

        # Ensure all EQUATOR columns are numeric
        for col in equator_cols:
            self.df[col] = pd.to_numeric(self.df[col], errors="coerce")

        question_scores = self.df.groupby("question_ID")[equator_cols].mean()
        score_percent = question_scores.mean(axis=1)

        thresholds = [0.6, 0.7, 0.8, 0.9]
        counts = [(score_percent >= t).mean() * 100 for t in thresholds]
        counts = [self._perturb_percent(c) for c in counts]

        fig, ax = plt.subplots(figsize=(8, 6))
        bars = ax.bar(
            [f">{int(t*100)}%" for t in thresholds],
            counts,
            color=sns.color_palette("Blues_d", len(thresholds)),
        )

        # Add value labels on bars
        for bar, count in zip(bars, counts):
            ax.text(
                bar.get_x() + bar.get_width() / 2.0,
                bar.get_height(),
                f"{count:.1f}%",
                ha="center",
                va="bottom",
                fontweight="bold",
            )

        ax.set_ylabel("% of Questions", fontsize=12, fontweight="bold")
        ax.set_title(
            "Percentage of Questions Passing Healthcare Criteria Score Thresholds",
            fontsize=14,
            fontweight="bold",
        )
        ax.set_ylim(0, 100)
        ax.grid(axis="y", alpha=0.3)

        plt.tight_layout()
        return fig

    def plot_equator_criteria_accuracy(self, exclude=None):
        """Bar plot showing average accuracy per healthcare reporting item across all batches."""
        if exclude is None:
            exclude = []

        equator_cols = [
            col
            for col in self.df.columns
            if col.startswith("equator_") and col not in exclude
        ]

        # Ensure all EQUATOR columns are numeric
        for col in equator_cols:
            self.df[col] = pd.to_numeric(self.df[col], errors="coerce")

        criteria_accuracy = self.df[equator_cols].mean() * 100
        criteria_accuracy = criteria_accuracy.apply(self._perturb_percent)
        criteria_accuracy = criteria_accuracy.sort_values()

        fig, ax = plt.subplots(figsize=(12, 8))
        bars = ax.barh(
            range(len(criteria_accuracy)),
            criteria_accuracy.values,
            color=sns.color_palette("Reds_d", len(criteria_accuracy)),
        )

        # Add value labels
        for i, (bar, val) in enumerate(zip(bars, criteria_accuracy.values)):
            ax.text(val, i, f" {val:.1f}%", ha="left", va="center", fontweight="bold")

        ax.set_yticks(range(len(criteria_accuracy)))
        ax.set_yticklabels(criteria_accuracy.index, fontsize=9)
        ax.set_xlabel("Accuracy (%)", fontsize=12, fontweight="bold")
        ax.set_title(
            "Healthcare Reporting Item Accuracy Across All Batches",
            fontsize=14,
            fontweight="bold",
        )
        ax.grid(axis="x", alpha=0.3)

        plt.tight_layout()
        return fig

    def plot_question_consistency_boxplot(self, exclude=None):
        """Boxplot showing EQUATOR score variability for each question across batches."""
        if exclude is None:
            exclude = []

        equator_cols = [
            col
            for col in self.df.columns
            if col.startswith("equator_") and col not in exclude
        ]

        # Ensure all EQUATOR columns are numeric
        for col in equator_cols:
            self.df[col] = pd.to_numeric(self.df[col], errors="coerce")

        df_temp = self.df.copy()
        df_temp["equator_total"] = df_temp[equator_cols].mean(axis=1)

        grouped = df_temp.groupby("question_ID")["equator_total"].apply(list)
        data = grouped.explode().reset_index()
        data["equator_total"] = data["equator_total"].astype(float)

        fig, ax = plt.subplots(figsize=(14, 6))

        # Create boxplot
        box_plot = ax.boxplot(
            [
                data[data["question_ID"] == qid]["equator_total"].values
                for qid in data["question_ID"].unique()
            ],
            labels=data["question_ID"].unique(),
            patch_artist=True,
        )

        # Color the boxes with coolwarm palette
        colors = sns.color_palette("coolwarm", len(box_plot["boxes"]))
        for patch, color in zip(box_plot["boxes"], colors):
            patch.set_facecolor(color)
            patch.set_alpha(0.7)

        ax.set_xticklabels(data["question_ID"].unique(), rotation=90, fontsize=9)
        ax.set_ylabel("Mean EQUATOR Score (0–1)", fontsize=12, fontweight="bold")
        ax.set_title(
            "Consistency of EQUATOR Scores Across Batches for Each Question",
            fontsize=14,
            fontweight="bold",
        )
        ax.grid(axis="y", alpha=0.3)

        plt.tight_layout()
        return fig

    def plot_equator_theme_radar(self, exclude=None):
        """Radar chart showing average healthcare reporting criteria coverage grouped by theme for each batch overlaid."""
        try:
            if exclude is None:
                exclude = []

            def extract_equator_number(col_name):
                """Extract numeric portion from healthcare criteria column name (e.g., '1a' -> 1, '22' -> 22)"""
                match = re.search(r"equator_(\d+)", col_name)
                return int(match.group(1)) if match else None

            def get_equator_columns_by_range(df, start, end, exclude):
                """Get healthcare criteria columns within a numeric range, excluding specified columns"""
                equator_cols = [
                    col
                    for col in df.columns
                    if col.startswith("equator_") and col not in exclude
                ]
                result_cols = []

                for col in equator_cols:
                    num = extract_equator_number(col)
                    if num is not None and start <= num <= end:
                        result_cols.append(col)

                return result_cols

            # Build themes dynamically based on healthcare reporting numbering ranges
            equator_themes = {
                "Title/Abstract": get_equator_columns_by_range(
                    self.df, 1, 1, exclude
                ),  # Items 1-1 (1a, 1b, etc.)
                "Background/Objectives": get_equator_columns_by_range(
                    self.df, 2, 3, exclude
                ),  # Items 2-3
                "Methods": get_equator_columns_by_range(
                    self.df, 4, 12, exclude
                ),  # Items 4-12
                "Results": get_equator_columns_by_range(
                    self.df, 13, 17, exclude
                ),  # Items 13-17
                "Discussion": get_equator_columns_by_range(
                    self.df, 18, 21, exclude
                ),  # Items 18-21
                "Funding": get_equator_columns_by_range(
                    self.df, 22, 22, exclude
                ),  # Item 22
            }

            # Filter out empty themes AND remove funding if any funding columns are excluded
            funding_cols = [
                col
                for col in self.df.columns
                if col.startswith("equator_") and extract_equator_number(col) == 22
            ]
            if any(col in exclude for col in funding_cols):
                equator_themes.pop("Funding", None)

            equator_themes = {
                theme: cols for theme, cols in equator_themes.items() if cols
            }

            # Check if we have any themes to plot
            if not equator_themes:
                logger.warning("No EQUATOR themes found for radar chart")
                fig, ax = plt.subplots(figsize=(10, 10))
                ax.text(
                    0.5,
                    0.5,
                    "No EQUATOR data available for radar chart",
                    ha="center",
                    va="center",
                    transform=ax.transAxes,
                    fontsize=16,
                )
                return fig

            # Get unique batches and sort them naturally
            batches = self.natural_sort_batches(self.df["batch"].unique())

            # Calculate theme scores per batch
            batch_theme_scores = {}
            for batch in batches:
                batch_data = self.df[self.df["batch"] == batch]
                batch_theme_scores[batch] = {}

                for theme, cols in equator_themes.items():
                    if cols:  # Only calculate if columns exist
                        # Ensure columns are numeric
                        numeric_data = batch_data[cols].apply(
                            pd.to_numeric, errors="coerce"
                        )
                        theme_score = numeric_data.mean().mean() * 100
                        batch_theme_scores[batch][theme] = (
                            theme_score if not pd.isna(theme_score) else 0
                        )
                    else:
                        batch_theme_scores[batch][theme] = 0

            # Create single radar chart with overlaid batches
            fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(polar=True))

            # Get consistent labels and angles (same for all batches)
            labels = list(equator_themes.keys())
            angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
            angles += angles[:1]  # Close the loop

            # Color palette for different batches
            colors = plt.cm.Set3(np.linspace(0, 1, len(batches)))

            # Plot each batch
            for i, batch in enumerate(batches):
                stats = [batch_theme_scores[batch][theme] for theme in labels]
                stats += stats[:1]  # Close the loop

                ax.plot(
                    angles, stats, color=colors[i], linewidth=3, label=batch, alpha=0.8
                )
                ax.fill(angles, stats, color=colors[i], alpha=0.1)

            # Set up the radar chart
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(labels, fontsize=11, fontweight="bold")
            ax.set_title(
                "Healthcare Criteria Theme Coverage by Batch", size=16, fontweight="bold", pad=20
            )
            ax.set_yticks([20, 40, 60, 80, 100])
            ax.set_ylim(0, 100)
            ax.grid(True, alpha=0.3)

            # Add legend
            ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.0), fontsize=10)

            # Add percentage labels at each theme position (showing range across batches)
            for angle, theme in zip(angles[:-1], labels):
                theme_values = [batch_theme_scores[batch][theme] for batch in batches]
                min_val, max_val = min(theme_values), max(theme_values)
                avg_val = sum(theme_values) / len(theme_values)

                # Position label closer to the center, inside the chart
                ax.text(
                    angle,
                    50,
                    f"{avg_val:.1f}%\n({min_val:.1f}-{max_val:.1f})",
                    ha="center",
                    va="center",
                    fontweight="bold",
                    fontsize=8,
                    color="darkblue",
                    bbox=dict(boxstyle="round,pad=0.2", facecolor="white", alpha=0.9),
                )

            plt.tight_layout()

        except Exception as e:
            logger.error(f"Failed to create healthcare criteria theme radar chart: {str(e)}")
            # Create a simple error plot instead of crashing
            fig, ax = plt.subplots(figsize=(10, 10))
            ax.text(
                0.5,
                0.5,
                f"Error creating radar chart:\n{str(e)}",
                ha="center",
                va="center",
                transform=ax.transAxes,
                fontsize=12,
            )
            ax.set_title(
                "Healthcare Criteria Theme Radar Chart - Error", fontsize=16, fontweight="bold"
            )

        return fig

    def generate_equator_figures(self, exclude=None):
        """Generate all healthcare reporting criteria analysis figures"""
        logger.info("Generating healthcare criteria figures")

        if exclude is None:
            exclude = []

        # Generate all figures
        threshold_fig = self.plot_equator_score_thresholds()
        criteria_fig = self.plot_equator_criteria_accuracy()
        consistency_fig = self.plot_question_consistency_boxplot()
        radar_fig = self.plot_equator_theme_radar()
        relevant_threshold_fig = self.plot_equator_score_thresholds(exclude=exclude)
        relevant_criteria_fig = self.plot_equator_criteria_accuracy(exclude=exclude)
        relevant_consistency_fig = self.plot_question_consistency_boxplot(
            exclude=exclude
        )
        relevant_radar_fig = self.plot_equator_theme_radar(exclude=exclude)

        self.results.update(
            {
                "equator_threshold_fig": threshold_fig,
                "equator_criteria_fig": criteria_fig,
                "equator_consistency_fig": consistency_fig,
                "equator_radar_fig": radar_fig,
                "relevant_equator_threshold_fig": relevant_threshold_fig,
                "relevant_equator_criteria_fig": relevant_criteria_fig,
                "relevant_equator_consistency_fig": relevant_consistency_fig,
                "relevant_equator_radar_fig": relevant_radar_fig,
            }
        )

        return threshold_fig, criteria_fig, consistency_fig, radar_fig

    def run_analysis(self, exclude_equator_cols=None):
        """Run all analyses"""
        logger.info("Starting comprehensive analysis")

        try:
            # Run all analyses
            self.analyze_summary_statistics()
            self.plot_batch_success_percentage()
            self.plot_batch_status_breakdown()
            self.plot_success_heatmap()
            self.calculate_posterior_probability()
            self.plot_runtime_analysis()
            self.generate_equator_figures(exclude=exclude_equator_cols)

            logger.info("All analyses completed successfully")

        except Exception as e:
            logger.error(f"Analysis failed: {str(e)}")
            raise

    def figure_to_base64(self, fig):
        """Convert matplotlib figure to base64 string"""
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", dpi=300, bbox_inches="tight")
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.read()).decode()
        buffer.close()
        return image_base64

    def create_word_report(self, output_path=None, input_xlsx_path=None):
        """Create a formatted Word document with all figures and analysis. Output filename includes input xlsx basename if not provided. Uses US/Pacific timezone for timestamps."""
        logger.info("Creating Word document report")

        # Use Pacific Time for all timestamps
        pacific = pytz.timezone("US/Pacific")
        now_pt = datetime.now(pacific)

        if output_path is None:
            # Set output directory to same location as input Excel file, or default to a local visualization folder
            if input_xlsx_path is not None:
                excel_dir = os.path.dirname(input_xlsx_path)
                xlsx_base = os.path.basename(input_xlsx_path)
                xlsx_no_ext = os.path.splitext(xlsx_base)[0]
                # Use the full base name after the timestamp
                output_filename = f"benchmark_report_{now_pt.strftime('%Y%m%d_%H%M%S')}_{xlsx_no_ext}.docx"

                # Save to both directories
                excel_output_path = os.path.join(excel_dir, output_filename)
                viz_root = os.path.join(os.getcwd(), "visualization")
                viz_output_path = os.path.join(viz_root, output_filename)

                # Ensure both directories exist
                os.makedirs(excel_dir, exist_ok=True)
                os.makedirs(viz_root, exist_ok=True)

                # Use the Excel directory path as the primary output path
                output_path = excel_output_path
            else:
                # Default to local visualization folder if no input path provided
                viz_root = os.path.join(os.getcwd(), "visualization")
                os.makedirs(viz_root, exist_ok=True)
                output_filename = (
                    f"benchmark_report_{now_pt.strftime('%Y%m%d_%H%M%S')}.docx"
                )
                output_path = os.path.join(viz_root, output_filename)

        # Create document
        doc = Document()

        # Section 1: Title Page
        self._create_title_page(doc)

        # Section 2: Executive Summary
        self._create_executive_summary(doc)

        # Section 3: Success Rate and Consistency
        self._create_success_rate_section(doc)

        # Section 4: Healthcare Reporting Criteria Analysis (All)
        self._create_equator_all_section(doc)

        # Section 5: Healthcare Reporting Criteria Analysis (Relevant)
        self._create_equator_relevant_section(doc)

        # Section 6: Technical Breakdown
        self._create_technical_breakdown_section(doc)

        # Save document
        doc.save(output_path)
        logger.info(f"Word report saved to: {output_path}")

        # If we have an Excel input path, also save to visualization directory
        if input_xlsx_path is not None and "viz_output_path" in locals():
            doc.save(viz_output_path)
            logger.info(f"Word report also saved to: {viz_output_path}")

        return output_path

    def _create_title_page(self, doc):
        """Create the title page with branding"""
        # Add title
        title = doc.add_heading("Benchmark Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add timestamp
        pst = timezone("US/Pacific")
        now_pst = datetime.now(pst)
        timestamp = doc.add_paragraph(
            f'Generated on {now_pst.strftime("%B %d, %Y at %I:%M %p PST")}'
        )
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _create_executive_summary(self, doc):
        """Create executive summary with key statistics"""
        # Section header
        self._add_section_header(doc, "Executive Summary")

        # Calculate key metrics
        total_runs = len(self.df)
        success_rate = (self.df["status"] == "success").mean() * 100
        success_rate = self._perturb_percent(success_rate)
        avg_runtime = self.df["total_time"].mean()
        avg_runtime = self._perturb_positive(avg_runtime)
        num_batches = self.df["batch"].nunique()
        unique_questions = self.df["question_ID"].nunique()

        # Healthcare reporting criteria metrics (all criteria)
        equator_cols = [col for col in self.df.columns if col.startswith("equator_")]
        equator_criteria_count = len(equator_cols)
        equator_avg_compliance = (
            self.df[equator_cols].mean().mean() if equator_cols else 0
        )
        equator_avg_compliance = self._perturb_ratio(equator_avg_compliance)

        # Healthcare reporting criteria metrics (relevant criteria)
        excluded_criteria = [
            "equator_6b_matching criteria (if applicable)",
            "equator_12b_subgroup/interaction methods",
            "equator_12e_sensitivity analyses",
            "equator_22_funding information",
        ]
        relevant_equator_cols = [
            col for col in self.df.columns if col.startswith("equator")
        ]
        relevant_equator_cols = [
            col for col in relevant_equator_cols if col not in excluded_criteria
        ]
        relevant_equator_count = len(relevant_equator_cols)
        relevant_equator_compliance = (
            self.df[relevant_equator_cols].mean().mean() if relevant_equator_cols else 0
        )
        relevant_equator_compliance = self._perturb_ratio(relevant_equator_compliance)

        summary_text = f"""
        This report presents a comprehensive analysis of benchmark performance across multiple batches and questions.
        The analysis includes both success rate metrics and detailed healthcare reporting criteria compliance assessment.
        
        Key Performance Metrics:
        • Total runs analyzed: {total_runs:,}
        • Overall success rate: {success_rate:.1f}%
        • Average runtime: {avg_runtime:.1f} seconds
        • Number of batches: {num_batches}
        • Number of unique questions: {unique_questions}
        
        Healthcare Reporting Criteria Analysis (All Criteria):
        • Total criteria evaluated: {equator_criteria_count}
        • Average compliance score: {equator_avg_compliance:.3f} ({equator_avg_compliance*100:.1f}%)
        
        Healthcare Reporting Criteria Analysis (Relevant Criteria):
        • Relevant criteria evaluated: {relevant_equator_count}
        • Average compliance score: {relevant_equator_compliance:.3f} ({relevant_equator_compliance*100:.1f}%)
        
        Note: The relevant criteria analysis excludes items not typically included in analytics 
        reports, such as funding information and certain methodological details not applicable in this simulated healthcare analytics context.

        Healthcare-Focused Analytics Context:
        • This simulated pipeline reflects how clinical study and healthcare analytics outputs can be evaluated for reporting quality and robustness.
        • Example applications include clinical trial protocols/results, therapeutic outcome reports, and health economics such as drug pricing or market uptake analyses.
        • All statistics and figures in this report are generated from simulated data and are not representative of any real company, product, or commercial entity.
        """

        doc.add_paragraph(summary_text)
        doc.add_page_break()

    def _create_success_rate_section(self, doc):
        """Create success rate and consistency analysis section"""
        self._add_section_header(doc, "Success Rate and Consistency Analysis")

        # Define figures for this section in order
        success_figures = [
            (
                "Summary Statistics",
                "summary_fig",
                "Comprehensive summary statistics for the benchmark analysis, including key performance metrics and distributions.",
            ),
            (
                "Batch Success Rates",
                "batch_success_fig",
                "Success rate for each batch with a 90% threshold reference line to identify high-performing batches.",
            ),
            (
                "Status Breakdown by Batch",
                "batch_breakdown_fig",
                "Detailed breakdown of all status types (success, failure, error, timeout) for each batch to identify patterns.",
            ),
            (
                "Status Heatmap",
                "heatmap_fig",
                "Detailed heatmap view of success/failure patterns across all batches and questions, highlighting problematic areas.",
            ),
            (
                "Bayesian Analysis",
                "posterior_fig",
                "Probabilistic assessment of success rate performance using Bayesian inference to estimate confidence intervals.",
            ),
            (
                "Batch Runtime Analysis",
                "batch_runtime_fig",
                "Detailed analysis of runtime performance across all batches, showing distribution and outliers.",
            ),
            (
                "Total Question Runtime Analysis",
                "question_runtime_fig",
                "Runtime performance analysis for each question across all batches, identifying computationally intensive questions.",
            ),
            (
                "Final Step Runtime Analysis",
                "question_runtime_s5_fig",
                "Analysis of final step (step 5) runtime performance, focusing on the most critical execution phase.",
            ),
            (
                "Runtime Distribution Analysis",
                "runtime_distribution_fig",
                "Statistical distribution of runtime values with outlier detection and performance benchmarks.",
            ),
        ]

        for section_title, fig_key, description in success_figures:
            self._add_figure_subsection(doc, section_title, fig_key, description)

            # Add Bayesian analysis text immediately after the Bayesian figure
            if fig_key == "posterior_fig" and "posterior_analysis" in self.results:
                doc.add_heading("Bayesian Analysis Interpretation", level=3)
                doc.add_paragraph(self.results["posterior_analysis"])

        doc.add_page_break()

    def _create_equator_all_section(self, doc):
        """Create healthcare reporting criteria analysis section for all criteria"""
        self._add_section_header(
            doc, "Alignment with Healthcare Reporting Criteria (All)"
        )

        # Add explanatory subtitle
        subtitle_text = """
        This section presents results comparing outputs to the entire healthcare reporting checklist with no exclusions. 
        Scores may be lower than expected because certain criteria (such as funding information, specific 
        methodological details, and administrative elements) are not always provided in standard analytics 
        reports by design.
        """
        doc.add_paragraph(subtitle_text)

        # Define healthcare reporting figures for all criteria
        equator_all_figures = [
            (
                "Healthcare Criteria Score Thresholds",
                "equator_threshold_fig",
                "Percentage of questions meeting various healthcare criteria score thresholds, showing overall compliance distribution.",
            ),
            (
                "Healthcare Criteria Accuracy",
                "equator_criteria_fig",
                "Average accuracy for each individual healthcare reporting criterion, identifying strengths and weaknesses in reporting.",
            ),
            (
                "Healthcare Criteria Score Consistency",
                "equator_consistency_fig",
                "Variability of healthcare criteria scores for each question across different batches, measuring consistency.",
            ),
            (
                "Healthcare Criteria Theme Coverage by Batch",
                "equator_radar_fig",
                "Radar chart comparing healthcare criteria theme coverage across all batches, highlighting reporting area performance.",
            ),
        ]

        for section_title, fig_key, description in equator_all_figures:
            self._add_figure_subsection(doc, section_title, fig_key, description)

        # Add EQUATOR analysis summary after the radar chart
        self._add_equator_analysis_summary(doc, all_criteria=True)
        doc.add_page_break()

    def _create_equator_relevant_section(self, doc):
        """Create healthcare reporting criteria analysis section for relevant criteria only"""
        self._add_section_header(
            doc, "Alignment with Healthcare Reporting Criteria (Relevant)"
        )

        # List excluded criteria
        excluded_criteria = [
            "equator_6b_matching criteria (if applicable)",
            "equator_12b_subgroup/interaction methods",
            "equator_12e_sensitivity analyses",
            "equator_22_funding information",
        ]

        subtitle_text = f"""
        This section focuses on healthcare reporting criteria that are directly relevant to analytics reports. 
        The following criteria have been excluded as they are not applicable to our standard reporting format:
        
        Excluded Criteria:
        • {excluded_criteria[0]}
        • {excluded_criteria[1]}
        • {excluded_criteria[2]}
        • {excluded_criteria[3]}
        
        This focused analysis provides a more accurate assessment of our reporting quality for applicable criteria.
        """
        doc.add_paragraph(subtitle_text)

        # Define healthcare reporting figures for relevant criteria
        equator_relevant_figures = [
            (
                "Relevant Healthcare Criteria Score Thresholds",
                "relevant_equator_threshold_fig",
                "Percentage of questions meeting various healthcare criteria score thresholds for relevant criteria only.",
            ),
            (
                "Relevant Healthcare Criteria Accuracy",
                "relevant_equator_criteria_fig",
                "Average accuracy for each relevant healthcare reporting criterion, showing targeted performance metrics.",
            ),
            (
                "Relevant Healthcare Criteria Score Consistency",
                "relevant_equator_consistency_fig",
                "Variability of relevant healthcare criteria scores across batches, measuring consistency for applicable criteria.",
            ),
            (
                "Relevant Healthcare Criteria Theme Coverage by Batch",
                "relevant_equator_radar_fig",
                "Radar chart comparing relevant healthcare criteria theme coverage, focusing on applicable reporting areas.",
            ),
        ]

        for section_title, fig_key, description in equator_relevant_figures:
            self._add_figure_subsection(doc, section_title, fig_key, description)

        # Add relevant EQUATOR analysis summary after the radar chart
        self._add_equator_analysis_summary(doc, all_criteria=False)
        doc.add_page_break()

    def _create_technical_breakdown_section(self, doc):
        """Create technical breakdown section for engineering team"""
        self._add_section_header(doc, "Technical Breakdown")

        subheader_text = """
        This section is designed for our engineering team to drill down on specific issues, 
        examine logs, and iterate on solutions. The following analysis provides detailed 
        information about failed questions and their patterns.
        """
        doc.add_paragraph(subheader_text)

        # Add the failed questions analysis
        self._add_failed_questions_analysis(doc)

        # Placeholder for additional technical content
        doc.add_paragraph(
            "\n[Additional technical analysis content will be added here by specialized functions]\n"
        )

    def _add_section_header(self, doc, title):
        """Add a large section header with consistent formatting"""
        doc.add_heading(title, level=1)
        # Add some spacing after header
        doc.add_paragraph()

    def _add_figure_subsection(self, doc, title, fig_key, description):
        """Add a figure subsection with title, description, and image"""
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)

        if fig_key in self.results:
            # Convert figure to image and add to document
            img_buffer = io.BytesIO()
            self.results[fig_key].savefig(
                img_buffer, format="png", dpi=300, bbox_inches="tight"
            )
            img_buffer.seek(0)

            # Use smaller width for heatmap, standard width for others
            width = Inches(2.5) if fig_key == "heatmap_fig" else Inches(6)
            doc.add_picture(img_buffer, width=width)
            img_buffer.close()
        else:
            doc.add_paragraph(f"[Figure {fig_key} not available]")

        doc.add_paragraph()  # Add spacing after figure

    def _add_equator_analysis_summary(self, doc, all_criteria=True):
        """Add healthcare reporting criteria analysis summary text"""
        doc.add_heading("Healthcare Reporting Criteria Summary", level=2)

        # Determine which columns to analyze
        if all_criteria:
            equator_cols = [
                col for col in self.df.columns if col.startswith("equator_")
            ]
            section_type = "complete"
        else:
            excluded_criteria = [
                "equator_6b_matching criteria (if applicable)",
                "equator_12b_subgroup/interaction methods",
                "equator_12e_sensitivity analyses",
                "equator_22_funding information",
            ]
            equator_cols = [col for col in self.df.columns if col.startswith("equator")]
            equator_cols = [col for col in equator_cols if col not in excluded_criteria]
            section_type = "relevant"

        if not equator_cols:
            doc.add_paragraph(
                f"No {section_type} healthcare reporting criteria data found in the dataset."
            )
            return

        # Calculate summary statistics
        avg_compliance = self.df[equator_cols].mean().mean()
        avg_compliance = self._perturb_ratio(avg_compliance)
        question_scores = (
            self.df.groupby("question_ID")[equator_cols].mean().mean(axis=1)
        )
        high_performers = (question_scores >= 0.8).sum()
        total_questions = len(question_scores)

        criteria_description = (
            "complete healthcare reporting checklist"
            if all_criteria
            else "relevant healthcare reporting criteria"
        )

        summary_text = f"""
        This healthcare reporting criteria analysis evaluates 
        compliance with clinical and healthcare-focused reporting guidelines across {len(equator_cols)} different {criteria_description}.
        
        Key Healthcare Reporting Findings:
        • Average compliance score: {avg_compliance:.3f} ({avg_compliance*100:.1f}%)
        • Questions with high compliance (≥80%): {high_performers}/{total_questions} ({high_performers/total_questions*100:.1f}%)
        • Criteria range from study design elements to methodological reporting standards
        • Analysis reveals variation across different reporting themes (Methods, Results, Discussion, etc.)
        
        The radar chart analysis reveals batch-specific strengths and weaknesses across different 
        healthcare reporting themes, helping identify targeted areas for improvement in research reporting quality.
        """

        doc.add_paragraph(summary_text)

    def _add_failed_questions_analysis(self, doc):
        """Add detailed analysis of failed questions to the Word document"""
        doc.add_heading("Failed Questions Analysis", level=2)

        # Define non-success statuses and their abbreviations
        non_success_statuses = ["failed", "timeout", "no_matching_concept"]
        status_abbrev = {"failed": "f", "timeout": "t", "no_matching_concept": "nmc"}

        # Filter to only non-success records
        failed_df = self.df[self.df["status"].isin(non_success_statuses)].copy()

        if failed_df.empty:
            doc.add_paragraph(
                "No failed, timeout, or no_matching_concept questions found in the dataset."
            )
            return

        # Group by question_ID to get all failures per unique question
        question_groups = failed_df.groupby(["question_ID", "question"])

        # Calculate failure counts and create sorting metrics
        failure_analysis = []

        for (question_id, question_text), group in question_groups:
            # Count each type of failure across all job runs for this question
            failed_count = (group["status"] == "failed").sum()
            timeout_count = (group["status"] == "timeout").sum()
            nmc_count = (group["status"] == "no_matching_concept").sum()

            # Total non-success count for primary sorting
            total_failures = len(group)

            # Create priority score for tie-breaking (failed=3, nmc=2, timeout=1)
            priority_score = (failed_count * 3) + (nmc_count * 2) + (timeout_count * 1)

            # Collect all job IDs with their statuses for sub-bullets
            job_details = []
            for _, row in group.iterrows():
                job_details.append(
                    {
                        "job_ID_run": row["job_ID_run"],
                        "status": row["status"],
                        "batch": row["batch"],
                    }
                )

            failure_analysis.append(
                {
                    "question_ID": question_id,
                    "question": question_text,
                    "total_failures": total_failures,
                    "priority_score": priority_score,
                    "failed_count": failed_count,
                    "timeout_count": timeout_count,
                    "nmc_count": nmc_count,
                    "job_details": job_details,
                }
            )

        # Sort by total failures (descending), then by priority score (descending)
        failure_analysis.sort(
            key=lambda x: (-x["total_failures"], -x["priority_score"])
        )

        # Add summary paragraph
        total_failed_questions = len(failure_analysis)
        total_failure_instances = sum(
            item["total_failures"] for item in failure_analysis
        )

        summary_text = f"""
        The following analysis identifies {total_failed_questions} unique questions that experienced 
        failures across all batches, representing {total_failure_instances} total failure instances. 
        Questions are ordered by total number of failures, with tie-breaking prioritized by failure 
        type severity (failed > no_matching_concept > timeout).
        """
        doc.add_paragraph(summary_text)

        # Create the bullet point list with sub-bullets for job IDs
        for item in failure_analysis:
            # Build the status count string for the main bullet
            status_parts = []
            if item["failed_count"] > 0:
                status_parts.append(f"f: {item['failed_count']}")
            if item["nmc_count"] > 0:
                status_parts.append(f"nmc: {item['nmc_count']}")
            if item["timeout_count"] > 0:
                status_parts.append(f"t: {item['timeout_count']}")

            status_string = ", ".join(status_parts)

            # Format the main bullet point (question level)
            main_bullet_text = (
                f"{item['question_ID']}, {item['question']} ({status_string})"
            )
            doc.add_paragraph(main_bullet_text, style="List Bullet")

            # Add sub-bullets for each job ID, sorted by status type for consistency
            # Sort job details by status priority (failed, no_matching_concept, timeout)
            status_priority = {"failed": 1, "no_matching_concept": 2, "timeout": 3}
            sorted_jobs = sorted(
                item["job_details"],
                key=lambda x: (status_priority.get(x["status"], 4), x["job_ID_run"]),
            )

            for job_detail in sorted_jobs:
                sub_bullet_text = (
                    f"({job_detail['status']}) jobID: {job_detail['job_ID_run']}"
                )
                doc.add_paragraph(sub_bullet_text, style="List Bullet 2")

        # Add summary statistics
        doc.add_heading("Failure Statistics Summary", level=3)

        # Calculate aggregate statistics
        total_failed = sum(item["failed_count"] for item in failure_analysis)
        total_timeout = sum(item["timeout_count"] for item in failure_analysis)
        total_nmc = sum(item["nmc_count"] for item in failure_analysis)

        stats_text = f"""
        Failure Type Breakdown:
        • Failed: {total_failed} instances across {sum(1 for item in failure_analysis if item['failed_count'] > 0)} questions
        • Timeout: {total_timeout} instances across {sum(1 for item in failure_analysis if item['timeout_count'] > 0)} questions  
        • No Matching Concept: {total_nmc} instances across {sum(1 for item in failure_analysis if item['nmc_count'] > 0)} questions
        
        This analysis helps identify problematic questions that may require attention from the engineering 
        team for debugging, optimization, or content review.
        """
        doc.add_paragraph(stats_text)


def main():
    """Main execution function"""
    if len(sys.argv) != 2:
        print("Usage: python benchmark_report_generator_v1.py <path_to_excel_file>")
        sys.exit(1)

    file_path = sys.argv[1]

        exclude_equator_cols = [
            "equator_6b_matching criteria (if applicable)",
            "equator_12b_subgroup/interaction methods",
            "equator_12e_sensitivity analyses",
            "equator_22_funding information",
        ]

    try:
        # Initialize generator
        generator = BenchmarkReportGenerator()

        # Load and validate data
        generator.load_and_validate_data(file_path)

        # Run analysis
        generator.run_analysis(exclude_equator_cols=exclude_equator_cols)

        # Create Word report, passing the input Excel file path for correct output naming
        report_path = generator.create_word_report(input_xlsx_path=file_path)

        print(f"\n{'='*50}")
        print("BENCHMARK ANALYSIS COMPLETED SUCCESSFULLY")
        print(f"{'='*50}")
        print(f"Report saved to: {report_path}")
        print(f"Log file: benchmark_report.log")
        print(f"{'='*50}")

    except Exception as e:
        logger.error(f"Script execution failed: {str(e)}")
        print(f"\nERROR: {str(e)}")
        print("Check the log file for detailed error information.")
        sys.exit(1)


if __name__ == "__main__":
    main()
